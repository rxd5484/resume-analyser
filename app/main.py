import os
import re
import uuid
from typing import List, Optional, Literal, Dict, Any, Tuple

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from dotenv import load_dotenv
import pdfplumber
import docx

from pydantic import BaseModel, Field

# LangChain (optional at runtime if OPENAI_API_KEY is set + quota available)
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # may be None for demo fallback
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")

app = FastAPI(title="Resume ↔ JD Matcher (Evidence-Based Demo)")




from fastapi.middleware.cors import CORSMiddleware
import os

ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "*")

origins = [o.strip() for o in ALLOWED_ORIGINS.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if origins == ["*"] else origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Demo in-memory storage (multi-user-ish)
# NOTE: This clears when the server restarts (fine for demo)
RESUMES: Dict[str, Dict[str, Any]] = {}


# -----------------------------
# Parsing helpers
# -----------------------------
def bytes_to_filelike(b: bytes):
    import io
    return io.BytesIO(b)


def normalize_text(t: str) -> str:
    t = t.replace("\x00", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import io

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return normalize_text("\n".join(text_parts))


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc_ = docx.Document(bytes_to_filelike(file_bytes))
    parts = []
    for p in doc_.paragraphs:
        if p.text:
            parts.append(p.text)
    return normalize_text("\n".join(parts))


def detect_filetype(filename: str) -> str:
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    return "unknown"


def parse_resume_to_text(resume: UploadFile, file_bytes: bytes) -> str:
    ftype = detect_filetype(resume.filename or "")
    if ftype == "pdf":
        return extract_text_from_pdf(file_bytes)
    if ftype == "docx":
        return extract_text_from_docx(file_bytes)
    raise HTTPException(status_code=400, detail="Unsupported file type. Upload a PDF or DOCX resume.")


# -----------------------------
# Hard blocker detection (demo)
# -----------------------------
HARD_BLOCKER_PATTERNS = [
    r"will not offer work authorization sponsorship",
    r"must not require work authorization sponsorship",
    r"will not provide any assistance.*immigration sponsorship",
    r"will not.*sign.*cpt",
    r"will not.*sign.*opt",
    r"\bno cpt\b",
    r"\bno opt\b",
    r"\bno sponsorship\b",
]


def detect_hard_blockers(jd_text: str) -> List[str]:
    jd = (jd_text or "").lower()
    hits = []
    for pat in HARD_BLOCKER_PATTERNS:
        if re.search(pat, jd):
            hits.append(pat)
    if hits:
        return ["Posting indicates NO sponsorship and/or NO CPT/OPT support."]
    return []


# -----------------------------
# Pydantic models (LLM outputs)
# -----------------------------
class Requirements(BaseModel):
    must_have: List[str] = Field(default_factory=list)
    nice_to_have: List[str] = Field(default_factory=list)


class EvidenceMatch(BaseModel):
    requirement: str
    status: Literal["covered", "not_found"]
    quote: Optional[str] = None
    chunk_id: Optional[int] = None
    confidence: float = 0.0
    reason: str = ""


# -----------------------------
# LangChain setup (optional usage)
# -----------------------------
req_parser = PydanticOutputParser(pydantic_object=Requirements)
match_parser = PydanticOutputParser(pydantic_object=EvidenceMatch)

requirements_prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You extract job requirements from a job description.\n"
            "Return JSON only that matches the schema.\n"
            "Rules:\n"
            "- must_have: core requirements, minimum qualifications\n"
            "- nice_to_have: preferred or bonus qualifications\n"
            "- Each requirement should be short, specific, and testable.\n"
            "- Do NOT include duplicates.\n"
            "{format_instructions}",
        ),
        ("human", "JOB DESCRIPTION:\n{jd_text}"),
    ]
)

verify_prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You verify whether a resume chunk satisfies a requirement.\n"
            "Return JSON only matching the schema.\n"
            "Hard rules:\n"
            "- If covered, quote MUST be an exact substring from the resume chunk.\n"
            "- If you cannot find exact evidence, status must be not_found and quote=null.\n"
            "- confidence must be between 0 and 1.\n"
            "{format_instructions}",
        ),
        (
            "human",
            "REQUIREMENT:\n{requirement}\n\n"
            "RESUME CHUNK (id={chunk_id}):\n{chunk_text}",
        ),
    ]
)


def get_llm() -> Optional[ChatOpenAI]:
    if not OPENAI_API_KEY:
        return None
    try:
        return ChatOpenAI(model=MODEL, temperature=0.0)
    except Exception:
        return None


def get_embeddings() -> Optional[OpenAIEmbeddings]:
    if not OPENAI_API_KEY:
        return None
    try:
        return OpenAIEmbeddings(model=EMBED_MODEL)
    except Exception:
        return None


# -----------------------------
# Utilities
# -----------------------------
def dedupe_clean(items: List[str]) -> List[str]:
    seen = set()
    cleaned = []
    for x in items:
        x = re.sub(r"\s+", " ", (x or "").strip())
        if not x:
            continue
        key = x.lower()
        if key not in seen:
            cleaned.append(x)
            seen.add(key)
    return cleaned


def split_resume(resume_text: str) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    chunks = splitter.split_text(resume_text)
    return chunks or []


def build_vectorstore(resume_text: str) -> Dict[str, Any]:
    """
    Tries to build FAISS with OpenAI embeddings.
    If quota/key issues happen, caller should fallback to local retrieval.
    """
    chunks = split_resume(resume_text)
    if not chunks:
        raise ValueError("No text chunks created from resume.")

    embeddings = get_embeddings()
    if embeddings is None:
        raise RuntimeError("Embeddings not available (missing OPENAI_API_KEY).")

    docs = [Document(page_content=c, metadata={"chunk_id": i}) for i, c in enumerate(chunks)]
    vs = FAISS.from_documents(docs, embeddings)
    return {"vectorstore": vs, "chunks": chunks}


# -----------------------------
# Local fallback retrieval + verification (no OpenAI)
# -----------------------------
def _tokens(s: str) -> List[str]:
    return re.findall(r"[a-zA-Z0-9\+\#\.]{2,}", (s or "").lower())


def _overlap_score(a: str, b: str) -> float:
    ta = set(_tokens(a))
    tb = set(_tokens(b))
    if not ta or not tb:
        return 0.0
    inter = len(ta & tb)
    denom = (len(ta) ** 0.5) * (len(tb) ** 0.5)
    return inter / denom if denom else 0.0


def local_top_chunks(requirement: str, chunks: List[str], top_k: int = 3) -> List[Tuple[int, str, float]]:
    scored = []
    for i, c in enumerate(chunks):
        scored.append((i, c, _overlap_score(requirement, c)))
    scored.sort(key=lambda x: x[2], reverse=True)
    return scored[:top_k]


def local_verify(requirement: str, chunk_id: int, chunk_text: str) -> EvidenceMatch:
    """
    Simple heuristic:
    - If overlap score is decent and we can extract an EXACT substring as quote, mark covered.
    - Quote is taken from the chunk itself (exact substring).
    """
    score = _overlap_score(requirement, chunk_text)
    if score < 0.18:
        return EvidenceMatch(
            requirement=requirement,
            status="not_found",
            quote=None,
            chunk_id=None,
            confidence=0.0,
            reason="Local fallback: low keyword overlap.",
        )

    # pick a short quote: first line/sentence that contains any token from requirement
    req_toks = set(_tokens(requirement))
    lines = [l.strip() for l in re.split(r"[\n\r]+", chunk_text) if l.strip()]
    best_line = None
    best_line_score = 0.0
    for line in lines[:20]:
        lt = set(_tokens(line))
        if not lt:
            continue
        s = len(req_toks & lt)
        if s > best_line_score:
            best_line_score = s
            best_line = line

    if not best_line:
        return EvidenceMatch(
            requirement=requirement,
            status="not_found",
            quote=None,
            chunk_id=None,
            confidence=0.0,
            reason="Local fallback: no good evidence line found.",
        )

    # ensure quote is exact substring
    quote = best_line
    if quote not in chunk_text:
        return EvidenceMatch(
            requirement=requirement,
            status="not_found",
            quote=None,
            chunk_id=None,
            confidence=0.0,
            reason="Local fallback: could not produce exact substring quote.",
        )

    conf = min(0.85, 0.4 + score)  # heuristic confidence
    return EvidenceMatch(
        requirement=requirement,
        status="covered",
        quote=quote,
        chunk_id=int(chunk_id),
        confidence=float(round(conf, 2)),
        reason="Local fallback: keyword overlap + exact substring evidence line.",
    )


# -----------------------------
# Requirement extraction
# -----------------------------
def extract_requirements_llm(jd_text: str) -> Requirements:
    llm = get_llm()
    if llm is None:
        raise RuntimeError("LLM not available (missing OPENAI_API_KEY).")
    chain = requirements_prompt | llm | req_parser
    out: Requirements = chain.invoke(
        {"jd_text": jd_text, "format_instructions": req_parser.get_format_instructions()}
    )
    out.must_have = dedupe_clean(out.must_have)
    out.nice_to_have = dedupe_clean(out.nice_to_have)
    return out


def extract_requirements_fallback(jd_text: str) -> Requirements:
    # Simple bullet/line fallback
    lines = [l.strip("•- \t") for l in (jd_text or "").splitlines() if len(l.strip()) > 6]
    lines = dedupe_clean(lines)
    # naive split
    return Requirements(must_have=lines[:10], nice_to_have=lines[10:20])


def extract_requirements(jd_text: str) -> Requirements:
    try:
        return extract_requirements_llm(jd_text)
    except Exception:
        return extract_requirements_fallback(jd_text)


# -----------------------------
# Verify matches (LLM if possible, else local)
# -----------------------------
def verify_requirement_against_top_chunks_llm(requirement: str, vs: FAISS, top_k: int = 3) -> EvidenceMatch:
    llm = get_llm()
    if llm is None:
        raise RuntimeError("LLM not available")

    docs = vs.similarity_search(requirement, k=top_k)
    best: Optional[EvidenceMatch] = None

    for d in docs:
        chunk_text = d.page_content
        chunk_id = int(d.metadata.get("chunk_id", -1))
        chain = verify_prompt | llm | match_parser
        candidate: EvidenceMatch = chain.invoke(
            {
                "requirement": requirement,
                "chunk_id": chunk_id,
                "chunk_text": chunk_text,
                "format_instructions": match_parser.get_format_instructions(),
            }
        )

        # enforce evidence substring rule
        if candidate.status == "covered":
            if not candidate.quote or candidate.quote not in chunk_text:
                candidate.status = "not_found"
                candidate.quote = None
                candidate.chunk_id = None
                candidate.confidence = 0.0
                candidate.reason = "Rejected: quote was not an exact substring of the resume chunk."

        if best is None:
            best = candidate
        else:
            if best.status != "covered" and candidate.status == "covered":
                best = candidate
            elif best.status == candidate.status and candidate.confidence > best.confidence:
                best = candidate

    if best is None:
        return EvidenceMatch(
            requirement=requirement,
            status="not_found",
            quote=None,
            chunk_id=None,
            confidence=0.0,
            reason="No chunks retrieved.",
        )
    return best


def verify_requirement_against_top_chunks_local(requirement: str, chunks: List[str], top_k: int = 3) -> EvidenceMatch:
    candidates = local_top_chunks(requirement, chunks, top_k=top_k)
    best: Optional[EvidenceMatch] = None
    for chunk_id, chunk_text, _ in candidates:
        cand = local_verify(requirement, chunk_id, chunk_text)
        if best is None:
            best = cand
        else:
            if best.status != "covered" and cand.status == "covered":
                best = cand
            elif best.status == cand.status and cand.confidence > best.confidence:
                best = cand
    if best is None:
        return EvidenceMatch(
            requirement=requirement,
            status="not_found",
            quote=None,
            chunk_id=None,
            confidence=0.0,
            reason="Local fallback: no chunks.",
        )
    return best


def compute_fit_score(matches_must: List[EvidenceMatch], matches_nice: List[EvidenceMatch]) -> float:
    def coverage(ms: List[EvidenceMatch]) -> float:
        if not ms:
            return 0.0
        covered = sum(1 for m in ms if m.status == "covered")
        return covered / len(ms)

    must_cov = coverage(matches_must)
    nice_cov = coverage(matches_nice)
    score = (0.75 * must_cov + 0.25 * nice_cov) * 100.0
    return round(score, 1)


# -----------------------------
# API models
# -----------------------------
class UploadResponse(BaseModel):
    resume_id: str
    chars: int


class AnalyzeRequest(BaseModel):
    resume_id: str
    jd_text: str


class AnalyzeResponse(BaseModel):
    fit_score: float
    hard_blockers: List[str] = Field(default_factory=list)
    method: str
    must_have: List[EvidenceMatch]
    nice_to_have: List[EvidenceMatch]
    missing_must_have: List[str]
    missing_nice_to_have: List[str]


# -----------------------------
# Routes
# -----------------------------
@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/upload", response_model=UploadResponse)
async def upload(resume: UploadFile = File(...)):
    file_bytes = await resume.read()
    resume_text = parse_resume_to_text(resume, file_bytes)

    if len(resume_text) < 200:
        raise HTTPException(
            status_code=400,
            detail="Resume text extraction seems too short. Try a different file or a text-based PDF.",
        )

    resume_id = str(uuid.uuid4())
    chunks = split_resume(resume_text)

    RESUMES[resume_id] = {
        "filename": resume.filename,
        "text": resume_text,
        "chunks": chunks,
    }

    return UploadResponse(resume_id=resume_id, chars=len(resume_text))


@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze(
    resume_id: str = Form(...),
    jd_text: str = Form(...),
):
    if not resume_id or resume_id not in RESUMES:
        raise HTTPException(status_code=404, detail="resume_id not found. Upload a resume first.")

    if len((jd_text or "").strip()) < 80:
        raise HTTPException(status_code=400, detail="Job description looks too short. Paste the full JD.")

    resume_text = RESUMES[resume_id]["text"]
    chunks: List[str] = RESUMES[resume_id]["chunks"]

    hard_blockers = detect_hard_blockers(jd_text)

    # 1) Try building vectorstore (OpenAI embeddings). If it fails (quota/key), fallback to local retrieval.
    vs = None
    method_parts = []
    try:
        built = build_vectorstore(resume_text)
        vs = built["vectorstore"]
        method_parts.append("faiss_openai_embeddings")
    except Exception as e:
        # quota / key / network -> local fallback
        vs = None
        method_parts.append(f"local_retrieval_fallback ({type(e).__name__})")

    # 2) Extract requirements (LLM if possible else fallback)
    reqs = extract_requirements(jd_text)

    # 3) Verify matches (LLM if possible + vs else local)
    must_matches: List[EvidenceMatch] = []
    nice_matches: List[EvidenceMatch] = []

    if vs is not None and get_llm() is not None:
        method_parts.append("llm_verify")
        for r in reqs.must_have:
            try:
                must_matches.append(verify_requirement_against_top_chunks_llm(r, vs, top_k=3))
            except Exception:
                # if LLM blows up mid-run (quota), degrade to local for that requirement
                must_matches.append(verify_requirement_against_top_chunks_local(r, chunks, top_k=3))
        for r in reqs.nice_to_have:
            try:
                nice_matches.append(verify_requirement_against_top_chunks_llm(r, vs, top_k=3))
            except Exception:
                nice_matches.append(verify_requirement_against_top_chunks_local(r, chunks, top_k=3))
    else:
        method_parts.append("local_verify")
        must_matches = [verify_requirement_against_top_chunks_local(r, chunks, top_k=3) for r in reqs.must_have]
        nice_matches = [verify_requirement_against_top_chunks_local(r, chunks, top_k=3) for r in reqs.nice_to_have]

    score = compute_fit_score(must_matches, nice_matches)

    # If hard blocker, cap score (demo realism)
    if hard_blockers:
        score = min(score, 40.0)

    missing_must = [m.requirement for m in must_matches if m.status != "covered"]
    missing_nice = [m.requirement for m in nice_matches if m.status != "covered"]

    return AnalyzeResponse(
        fit_score=score,
        hard_blockers=hard_blockers,
        method=" + ".join(method_parts),
        must_have=must_matches,
        nice_to_have=nice_matches,
        missing_must_have=missing_must,
        missing_nice_to_have=missing_nice,
    )


# OPTIONAL: keep your old one-shot endpoint for convenience (upload+analyze in one request)
# This is handy for quick testing in /docs, but the dashboard should use /upload then /analyze.
@app.post("/analyze_one_shot", response_model=AnalyzeResponse)
async def analyze_one_shot(
    jd_text: str = Form(...),
    resume: UploadFile = File(...),
):
    up = await upload(resume)  # stores it, returns resume_id
    return await analyze(resume_id=up.resume_id, jd_text=jd_text)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app.main:app",
        host="0.0.0.0",
        port=int(os.getenv("PORT", 8000))
    )